import pandas as pd
import pytesseract
from PIL import Image
import fitz 
from docx import Document
import docx
from PIL import Image

def parse_file(file_path):
    if file_path.endswith('.csv'):
        return pd.read_csv(file_path)
    elif file_path.endswith('.xlsx'):
        return pd.read_excel(file_path)
    elif file_path.endswith('.txt'):
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_path.endswith('.docx'):
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith('.pdf'):
        doc = fitz.open(file_path)
        return "\n".join([page.get_text() for page in doc])
    elif file_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        image = Image.open(file_path)
        return pytesseract.image_to_string(image)
    else:
        return "Unsupported file type."
    
    
def extract_text_from_file(filepath):
    ext = filepath.split('.')[-1].lower()
    if ext == 'pdf':
        return extract_text_from_pdf(filepath)
    elif ext == 'docx':
        return extract_text_from_docx(filepath)
    elif ext in ['png', 'jpg', 'jpeg']:
        return extract_text_from_image(filepath)
    elif ext == 'txt':
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise Exception("Unsupported file type for text extraction.")

def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    return pytesseract.image_to_string(image)
